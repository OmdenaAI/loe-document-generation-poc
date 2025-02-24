import streamlit as st
import docx
import json
import re
from io import BytesIO
import openai
import os
from datetime import datetime

# Set OpenAI API key
openai.api_key = st.secrets.OPENAI_API_KEY

# Initialize session state
if "template_data" not in st.session_state:
    st.session_state.template_data = None
if "ai_suggestions" not in st.session_state:
    st.session_state.ai_suggestions = {}
if "accepted_ai_suggestions" not in st.session_state:
    st.session_state.accepted_ai_suggestions = {}
if "placeholders" not in st.session_state:
    st.session_state.placeholders = {}
if "processed_paragraphs" not in st.session_state:
    st.session_state.processed_paragraphs = []
if "form_data" not in st.session_state:
    st.session_state.form_data = {}
if "confirm_download" not in st.session_state:
    st.session_state.confirm_download = False
if "download_triggered" not in st.session_state:
    st.session_state.download_triggered = False
if "template_json" not in st.session_state:  # New: Store template.json in memory
    st.session_state.template_json = None

# ✅ Partner Tab Functions (unchanged except where noted)
def extract_paragraphs_from_docx(docx_file):
    """Extract paragraphs from DOCX with formatting, removing Markdown artifacts."""
    doc = docx.Document(docx_file)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            styled_text = ""
            for run in para.runs:
                text = run.text
                text = re.sub(r'```markdown\s*|\s*```', '', text)
                if run.bold:
                    text = f"<b>{text}</b>"
                if run.italic:
                    text = f"<i>{text}</i>"
                if run.underline:
                    text = f"<u>{text}</u>"
                styled_text += text
            paragraphs.append({"text": styled_text, "style": para.style.name})
    return paragraphs

def extract_placeholders(paragraphs):
    text = "\n".join([p["text"] for p in paragraphs])
    extracted = list(set(re.findall(r"\$\{(.*?)\}", text)))
    return {ph: {"type": "Text", "required": False, "is_conditional": False, "dependent_on": []} for ph in extracted}

def suggest_placeholders_with_ai(paragraphs):
    markdown_text = "\n\n".join([p["text"] for p in paragraphs])
    prompt = f"""
    Extract missing placeholders from this document. Respond with ONLY a JSON in this format:
    {{
      "suggested_placeholders": {{
        "field1": "descriptive-name-1",
        "field2": "descriptive-name-2"
      }}
    }}
    **Rules:** 
    - Return ONLY JSON, no explanations.
    - Each key should be lowercase with hyphens.
    **Document:**
    {markdown_text}
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI that must return JSON only. No extra text."},
                {"role": "user", "content": prompt}
            ]
        )
        raw_response = response.choices[0].message.content.strip()
        if not raw_response:
            st.error("⚠️ OpenAI API returned an empty response.")
            return {}
        try:
            parsed_response = json.loads(raw_response)
        except json.JSONDecodeError:
            st.error("⚠️ OpenAI returned malformed JSON.")
            return {}
        if "suggested_placeholders" not in parsed_response:
            st.error("⚠️ OpenAI response did not contain 'suggested_placeholders'.")
            return {}
        return {
            key: {
                "type": "Text",
                "required": False,
                "is_conditional": False,
                "dependent_on": []
            }
            for key in parsed_response["suggested_placeholders"].values()
        }
    except Exception as e:
        st.error(f"⚠️ OpenAI API error: {str(e)}")
        return {}

def insert_placeholders_in_markdown(paragraphs, placeholders):
    markdown_text = "\n\n".join([p["text"] for p in paragraphs])
    placeholder_list = "\n".join([f"- {p}" for p in placeholders.keys()])
    prompt = f"""
    Insert these placeholders into the document:
    {placeholder_list}
    **Rules:**
    1. Keep existing placeholders in their positions.
    2. Insert new placeholders logically.
    3. Preserve Markdown formatting.
    **Document:**
    {markdown_text}
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "Return the formatted Markdown document."},
                      {"role": "user", "content": prompt}]
        )
        updated_text = response.choices[0].message.content.strip().split("\n\n")
        return [{"text": re.sub(r'```markdown\s*|\s*```', '', para), "style": "Normal"} for para in updated_text]
    except Exception as e:
        st.error(f"⚠️ OpenAI API error: {str(e)}")
        return paragraphs

def convert_markdown_to_docx(paragraphs, output_path="updated_template.docx"):
    doc = docx.Document()
    for para in paragraphs:
        clean_text = re.sub(r'```markdown\s*|\s*```', '', para["text"])
        if para["style"].startswith("Heading"):
            level = int(para["style"].split()[-1])
            doc.add_paragraph(clean_text, style=f"Heading {min(level, 9)}")
        else:
            doc.add_paragraph(clean_text)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Reset buffer position to the beginning
    doc.save(output_path)
    return output_path

def render_partner_preview(paragraphs, placeholders, accepted_suggestions):
    if not paragraphs:
        return ["<p>No document uploaded yet.</p>"]
    preview_paragraphs = []
    base_style = "font-family: Arial, sans-serif; font-size: 14px; line-height: 1.5; margin-bottom: 10px;"
    for para in paragraphs:
        styled_para = re.sub(r'```markdown\s*|\s*```', '', para["text"])
        for ph, settings in placeholders.items():
            bg_color = "#e6f3ff" if not settings["required"] else "#ffe6e6"
            tooltip = f"Type: {settings['type']}, Required: {settings['required']}, Conditional: {settings['is_conditional']}"
            styled_ph = f'<span style="background-color: {bg_color}; padding: 2px 4px; border-radius: 3px;" title="{tooltip}">${{{ph}}}</span>'
            styled_para = styled_para.replace(f"${{{ph}}}", styled_ph)
        for ph in accepted_suggestions:
            if f"${{{ph}}}" not in para["text"]:
                styled_para += f' <span style="color: gray; font-style: italic;" title="AI-suggested (position TBD)">${{{ph}}}</span>'
        if para["style"].startswith("Heading"):
            level = int(para["style"].split()[-1])
            preview_paragraphs.append(f'<h{level} style="{base_style}">{styled_para}</h{level}>')
        else:
            preview_paragraphs.append(f'<div style="{base_style}">{styled_para}</div>')
    return preview_paragraphs

# ✅ User Tab Functions (unchanged except where noted)
def should_show_field(field_name, field_info, form_data):
    dependent_on = field_info.get("dependent_on")
    if not field_info.get("is_conditional", False) or not dependent_on:
        return True
    parent_field = dependent_on[0] if isinstance(dependent_on, list) else dependent_on
    parent_value = form_data.get(parent_field.strip("${}"), None)
    return parent_value not in [None, "", [], False]

def is_field_required(field_name, field_info, form_data):
    if field_info.get("required", False):
        if field_info.get("is_conditional", False):
            dependent_on = field_info.get("dependent_on")
            if dependent_on:
                parent_field = dependent_on[0] if isinstance(dependent_on, list) else dependent_on
                parent_value = form_data.get(parent_field.strip("${}"), None)
                return parent_value not in [None, "", [], False]
        return True
    return False

def validate_form(form_data, placeholders):
    missing_fields = [
        name.strip("${}") for name, info in placeholders.items()
        if is_field_required(name, info, form_data) and not form_data.get(name.strip("${}"))
    ]
    if missing_fields:
        st.error(f"⚠️ Please fill in all required fields: {', '.join(missing_fields)}")
        return False
    return True

def render_user_preview(doc_path, form_data):
    try:
        doc = docx.Document(doc_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    except FileNotFoundError:
        return ["<p>Please generate the template in the Partner tab first.</p>"]
    
    preview_paragraphs = []
    base_style = "font-family: Arial, sans-serif; font-size: 14px; line-height: 1.5; margin-bottom: 10px;"
    for para in paragraphs:
        styled_para = re.sub(r'```markdown\s*|\s*```', '', para)
        for field_name, value in form_data.items():
            placeholder = f"${{{field_name}}}"
            if placeholder in styled_para:
                if value and value not in ["", None]:
                    styled_para = styled_para.replace(placeholder, f'<span style="background-color: #e6ffe6; padding: 2px 4px; border-radius: 3px;">{value}</span>')
                else:
                    styled_para = styled_para.replace(placeholder, f'<span style="background-color: #fff3e6; padding: 2px 4px; border-radius: 3px;">${{{field_name}}}</span>')
        preview_paragraphs.append(f'<div style="{base_style}">{styled_para}</div>')
    return preview_paragraphs

def clean_up_document_with_llm(original_text, filled_data):
    prompt = f"""
    You are an AI that formats documents by replacing placeholders with provided data.
    Instructions:
    - Replace all placeholders (e.g., ${{field-name}}) with the given values.
    - If a placeholder is optional and left blank, remove any **sentence, bullet point, or paragraph** that depends on it.
    - If a placeholder is optional and left blank and there is a section or block that depends on it, remove the entire section or block.
    - Ensure the document remains **flowing naturally** without empty spaces or missing context.
    - Do not include Markdown syntax (e.g., ```, #, *) in the output; return plain text only.
    **Original Document:**
    {original_text}
    **Filled Data:**
    {json.dumps(filled_data, indent=4)}
    Return the **cleaned-up** document as plain text.
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are an expert legal document formatter."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"⚠️ LLM Error: {str(e)}")
        return original_text

# Main app with tabs
st.title("📄 Placeholder Management App")
partner_tab, user_tab = st.tabs(["Partner", "User"])

# ✅ Partner Tab Content
with partner_tab:
    st.header("Partner Dashboard")
    uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

    if uploaded_file and not st.session_state.processed_paragraphs:
        st.session_state.processed_paragraphs = extract_paragraphs_from_docx(uploaded_file)
        st.session_state.placeholders = extract_placeholders(st.session_state.processed_paragraphs)
        if not st.session_state.ai_suggestions:
            st.write("🤖 AI is analyzing the document...")
            st.session_state.ai_suggestions = suggest_placeholders_with_ai(st.session_state.processed_paragraphs)

    with st.expander("📑 Document Preview", expanded=True):
        all_placeholders = {**st.session_state.placeholders, **st.session_state.accepted_ai_suggestions}
        preview_paragraphs = render_partner_preview(st.session_state.processed_paragraphs, all_placeholders, st.session_state.accepted_ai_suggestions)
        st.markdown('<div style="max-height: 400px; overflow-y: auto;">', unsafe_allow_html=True)
        for para in preview_paragraphs:
            st.markdown(para, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("<small>Legend: Light blue = Optional, Light red = Required, Gray italic = AI-suggested</small>", unsafe_allow_html=True)

    st.markdown("### 📝 Existing Placeholders")
    for ph, settings in st.session_state.placeholders.items():
        st.markdown(f"#### {ph}")
        settings["type"] = st.selectbox("Type", ["Text", "Date", "Number"], key=f"dt_{ph}")
        settings["required"] = st.checkbox("📍 Required", key=f"req_{ph}")
        settings["is_conditional"] = st.checkbox("🔗 Conditional", key=f"cond_{ph}")
        if settings["is_conditional"]:
            all_placeholders_list = list(st.session_state.placeholders.keys()) + list(st.session_state.accepted_ai_suggestions.keys())
            all_placeholders_list = [p for p in all_placeholders_list if p != ph]
            settings["dependent_on"] = st.multiselect("Depends on:", all_placeholders_list, default=settings["dependent_on"], key=f"dep_{ph}")

    st.markdown("### 🧠 AI-Suggested Placeholders")
    for ph in st.session_state.ai_suggestions.keys():
        accept = st.checkbox(f"✅ Accept: {ph}", key=f"accept_{ph}")
        if accept:
            settings = {
                "type": st.selectbox("Type", ["Text", "Date", "Number"], key=f"dt_ai_{ph}"),
                "required": st.checkbox("📍 Required", key=f"req_ai_{ph}"),
                "is_conditional": st.checkbox("🔗 Conditional", key=f"cond_ai_{ph}"),
                "dependent_on": []
            }
            if settings["is_conditional"]:
                all_placeholders_list = list(st.session_state.placeholders.keys()) + list(st.session_state.accepted_ai_suggestions.keys())
                all_placeholders_list = [p for p in all_placeholders_list if p != ph]
                settings["dependent_on"] = st.multiselect("Depends on:", all_placeholders_list, default=settings["dependent_on"], key=f"dep_ai_{ph}")
            st.session_state.accepted_ai_suggestions[ph] = settings

    all_placeholders = {**st.session_state.placeholders, **st.session_state.accepted_ai_suggestions}

    if st.button("💾 Save Template"):
        all_placeholders = {**st.session_state.placeholders, **st.session_state.accepted_ai_suggestions}
        # Save to session_state instead of file
        st.session_state.template_json = {"placeholders": all_placeholders}
        st.success("✅ Template saved!")
        if st.session_state.accepted_ai_suggestions:
            st.write("🔹 AI placeholders accepted – updating document...")
            updated_paragraphs = insert_placeholders_in_markdown(st.session_state.processed_paragraphs, all_placeholders)
        else:
            st.write("✅ No AI placeholders accepted – using extracted placeholders only.")
            updated_paragraphs = st.session_state.processed_paragraphs
        if updated_paragraphs and isinstance(updated_paragraphs[0], str):
            updated_paragraphs = [{"text": para, "style": "Normal"} for para in updated_paragraphs]
        st.session_state.processed_paragraphs = updated_paragraphs

        # Provide download option for template.json
        json_bytes = BytesIO(json.dumps(st.session_state.template_json, indent=4).encode("utf-8"))
        st.download_button("📥 Download Template JSON", json_bytes, "template.json", "application/json")
        updated_doc_path = convert_markdown_to_docx(st.session_state.processed_paragraphs)
        with open(updated_doc_path, "rb") as doc_file:
            st.download_button("📥 Download Updated DOCX", doc_file, "updated_template.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ✅ User Tab Content
with user_tab:
    st.header("User Dashboard")
    st.title("📝 Dynamic Form & AI-Powered Document Generation")

    # Check for template in session_state instead of file
    if not st.session_state.template_json:
        st.error("❌ No template available. Please save a template in the Partner tab first.")
        st.stop()
    else:
        template_data = st.session_state.template_json

    placeholders = template_data.get("placeholders", {})
    if set(st.session_state.form_data.keys()) != set(key.strip("${}") for key in placeholders):
        st.session_state.form_data = {key.strip("${}"): None for key in placeholders}

    st.write("### ✏️ Fill out the form below:")
    form_data = {}
    for field_name_with_syntax, field_info in placeholders.items():
        field_name = field_name_with_syntax.strip("${}")
        if should_show_field(field_name, field_info, st.session_state.form_data):
            label = f"{field_name} {'*' if is_field_required(field_name, field_info, st.session_state.form_data) else ''}"
            if field_info["type"] == "Text":
                form_data[field_name] = st.text_input(label, key=f"input_{field_name}")
            elif field_info["type"] == "Number":
                form_data[field_name] = st.number_input(label, key=f"input_{field_name}")
            elif field_info["type"] == "Date":
                selected_date = st.date_input(label, key=f"input_{field_name}")
                form_data[field_name] = selected_date.strftime('%Y-%m-%d')
            else:
                form_data[field_name] = st.text_input(label, key=f"input_{field_name}")
            st.session_state.form_data[field_name] = form_data[field_name]

    with st.expander("📑 Document Preview", expanded=True):
        preview_paragraphs = render_user_preview("updated_template.docx", st.session_state.form_data)
        st.markdown('<div style="max-height: 400px; overflow-y: auto;">', unsafe_allow_html=True)
        for para in preview_paragraphs:
            st.markdown(para, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("<small>Legend: Green = Filled, Orange = Empty</small>", unsafe_allow_html=True)

    if st.button("📥 Download Final Document"):
        if validate_form(st.session_state.form_data, placeholders):
            empty_fields = [
                name.strip("${}") for name, value in st.session_state.form_data.items()
                if not value or value in ["", None]
            ]
            proceed = True
            if empty_fields:
                st.warning("The following placeholders will be left empty:")
                st.write(", ".join(empty_fields))
                st.write("Are you sure you want to proceed with these fields empty? This is a legal document, so please confirm your intent.")
                proceed = st.button("Confirm and Download")

            if proceed:
                try:
                    # Load the template document
                    doc = docx.Document("updated_template.docx")
                    original_text = "\n".join([para.text for para in doc.paragraphs])
                    cleaned_text = clean_up_document_with_llm(original_text, st.session_state.form_data)

                    # Create the final document in memory
                    final_doc = docx.Document()
                    for line in cleaned_text.split("\n"):
                        if line.strip():
                            final_doc.add_paragraph(line)

                    # Save to BytesIO instead of disk
                    buffer = BytesIO()
                    final_doc.save(buffer)
                    buffer.seek(0)

                    st.success("✅ AI-processed document generated successfully!")
                    st.download_button(
                        label="📥 Download AI-Cleaned Document",
                        data=buffer,
                        file_name="final_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except FileNotFoundError:
                    st.error("❌ 'updated_template.docx' not found. Please save the template in the Partner tab first.")
                except Exception as e:
                    st.error(f"❌ Error generating document: {str(e)}")