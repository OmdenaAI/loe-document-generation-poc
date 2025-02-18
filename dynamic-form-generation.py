import streamlit as st
import json
import os
import docx
import openai
from datetime import datetime

# Set OpenAI API Key
openai.api_key = st.secrets.OPENAI_API_KEY #os.getenv("OPENAI_API_KEY")

st.title("📝 Dynamic Form & AI-Powered Document Generation")

# ✅ Load template JSON
template_path = "template.json"
if not os.path.exists(template_path):
    st.error("❌ Missing `template.json`. Please complete the document upload step first.")
    st.stop()

with open(template_path, "r") as f:
    template_data = json.load(f)

placeholders = template_data.get("placeholders", {})

# ✅ Initialize session state for form data if not exists
if "form_data" not in st.session_state:
    st.session_state.form_data = {key.strip("${}"): None for key in placeholders}

# ✅ Function to check if a dependent field should be displayed
def should_show_field(field_name, field_info, form_data):
    """
    Determines whether a field should be displayed based on its dependencies.
    
    - Conditional fields appear **only if the parent field is filled**.
    - If a field has no dependency, it is always shown.
    """
    dependent_on = field_info.get("dependent_on")

    if not field_info.get("is_conditional", False) or not dependent_on:
        return True  # Always show non-conditional fields

    # Ensure correct reference to parent field
    parent_field = dependent_on.strip("${}")
    parent_value = form_data.get(parent_field, None)

    return parent_value not in [None, "", [], False]  # Show field **only when parent is filled**

# ✅ Function to check if a field should be required
def is_field_required(field_name, field_info, form_data):
    """
    Determines whether a field is required based on:
    - Whether it is marked as required.
    - Whether its parent field is filled (if it is conditional).
    """
    if field_info.get("required", False):
        if field_info.get("is_conditional", False):
            dependent_on = field_info.get("dependent_on")
            if dependent_on:
                parent_value = form_data.get(dependent_on.strip("${}"), None)
                return parent_value not in [None, "", [], False]  # Required **only when parent is filled**
        return True  # Standard required fields

    return False  # Default: Not required

# ✅ Dynamic Form UI
st.write("### ✏️ Fill out the form below:")
form_data = {}

for field_name_with_syntax, field_info in placeholders.items():
    field_name = field_name_with_syntax.strip("${}")  # Remove ${} syntax for internal use

    # ✅ Display only when dependency is filled
    if should_show_field(field_name, field_info, st.session_state.form_data):  
        label = f"{field_name} {'*' if is_field_required(field_name, field_info, st.session_state.form_data) else ''}"

        # ✅ Handle different field types
        if field_info["type"] == "Text":
            form_data[field_name] = st.text_input(label, key=f"input_{field_name}")
        elif field_info["type"] == "Number":
            form_data[field_name] = st.number_input(label, key=f"input_{field_name}")
        elif field_info["type"] == "Date":
            selected_date = st.date_input(label, key=f"input_{field_name}")
            form_data[field_name] = selected_date.strftime('%Y-%m-%d')
        elif field_info["type"] == "Dropdown":
            options = ["Option 1", "Option 2", "Option 3"]
            form_data[field_name] = st.selectbox(label, options, key=f"input_{field_name}")
        elif field_info["type"] == "Checkbox":
            form_data[field_name] = st.checkbox(label, key=f"input_{field_name}")
        else:
            form_data[field_name] = st.text_input(label, key=f"input_{field_name}")

        st.session_state.form_data[field_name] = form_data[field_name]

# ✅ Validation function to ignore hidden required fields
def validate_form(form_data, placeholders):
    """
    Ensures all required fields are filled **but ignores hidden fields**.
    - Conditional fields are **only required if their parent is filled**.
    """
    missing_fields = [
        name.strip("${}") for name, info in placeholders.items()
        if is_field_required(name, info, form_data) and not form_data.get(name.strip("${}"))
    ]
    
    if missing_fields:
        st.error(f"⚠️ Please fill in all required fields: {', '.join(missing_fields)}")
        return False
    return True

# ✅ Function to process document with LLM
def clean_up_document_with_llm(original_text, filled_data):
    """Uses LLM to remove sections where optional placeholders are missing."""
    
    prompt = f"""
    You are an AI that formats documents by replacing placeholders with provided data.
    
 Instructions:
    - Replace all placeholders (e.g., ${{field-name}}) with the given values.
    - If a placeholder is optional and left blank, remove any **sentence, bullet point, or paragraph** that depends on it.
    - If a placeholder is optional and left blank and there is a section or a block in the document that depends on it, remove the entire section or block.
     - Ensure the document remains **flowing naturally** without empty spaces or missing context.
     - For example, if a legal WILL document includes a section for "Children" and the user does not have any children, remove the entire "Children" section and so on.
    - Another for example, if a legal WILL document includes a section for "Gifts" and the user does not have any gift to distribute, remove the entire "Gift" section and so on.
    **Original Document:**

    {original_text}
    
    **Filled Data:**
    {json.dumps(filled_data, indent=4)}
    
    Return the **cleaned-up** document.
    """

    try:
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are an expert legal document formatter."},
                {"role": "user", "content": prompt}
            ]
        )
        
        return response.choices[0].message.content.strip()

    except Exception as e:
        st.error(f"⚠️ LLM Error: {str(e)}")
        return original_text  # Return the original text if LLM fails

# ✅ Document Generation Logic with LLM
if st.button("📥 Download Final Document"):
    if validate_form(form_data, placeholders):
        # Save filled data
        with open("filled_data.json", "w") as f:
            json.dump(form_data, f, indent=4)

        try:
            doc = docx.Document("updated_template.docx")
            original_text = "\n".join([para.text for para in doc.paragraphs])

            # Process document using LLM
            cleaned_text = clean_up_document_with_llm(original_text, form_data)

            # ✅ Create new document with cleaned-up text
            final_doc = docx.Document()
            for line in cleaned_text.split("\n"):
                if line.strip():
                    final_doc.add_paragraph(line)
            
            # Save the final document
            output_path = "final_document.docx"
            final_doc.save(output_path)
            
            st.success("✅ AI-processed document generated successfully!")

            # ✅ Download Button
            with open(output_path, "rb") as f:
                st.download_button(
                    "📥 Download AI-Cleaned Document",
                    f,
                    file_name="final_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"❌ Error generating document: {str(e)}")
